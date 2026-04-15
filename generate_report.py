"""
Generate the IEEE-style single-column project report (.docx).
Run: python generate_report.py
"""

import os
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement

# ── Paths ──
BASE_DIR = Path(__file__).parent
IMG_DIR = Path(r"C:\Users\prakh\.gemini\antigravity\brain\529b1c07-d187-4b29-a7d5-d775f69988ec")
OUTPUT_PATH = BASE_DIR / "SkinDisease_Project_Report.docx"

def find_image(prefix):
    for f in IMG_DIR.glob(f"{prefix}*.png"):
        return str(f)
    return None

IMG_ARCH = find_image("system_architecture")
IMG_PIPELINE = find_image("training_pipeline")
IMG_PREPROCESS = find_image("data_preprocessing")
IMG_INFERENCE = find_image("inference_pipeline")

ROMAN = {1:'I',2:'II',3:'III',4:'IV',5:'V',6:'VI',7:'VII',8:'VIII',9:'IX',10:'X'}

# ── Helpers ──
def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_padding(cell, top=30, bottom=30, start=40, end=40):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top',top),('bottom',bottom),('start',start),('end',end)]:
        m = OxmlElement(f'w:{side}')
        m.set(qn('w:w'), str(val))
        m.set(qn('w:type'), 'dxa')
        tcMar.append(m)
    tcPr.append(tcMar)

def body(doc, text, space_after=4, space_before=0, indent=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.first_line_indent = Cm(0.5) if indent else Cm(0)
    p.paragraph_format.line_spacing = Pt(13)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(10)
    return p

def section_heading(doc, num, title):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(f"{ROMAN.get(num, str(num))}. {title.upper()}")
    r.font.name = 'Times New Roman'
    r.font.size = Pt(11)
    r.bold = True

def subsection_heading(doc, letter, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(f"{letter}. {title}")
    r.font.name = 'Times New Roman'
    r.font.size = Pt(10)
    r.italic = True
    r.bold = True

def bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.7)
    p.paragraph_format.first_line_indent = Cm(-0.35)
    p.paragraph_format.line_spacing = Pt(13)
    br = p.add_run("\u2022 ")
    br.font.name = 'Times New Roman'
    br.font.size = Pt(10)
    if bold_prefix:
        rb = p.add_run(bold_prefix)
        rb.font.name = 'Times New Roman'
        rb.font.size = Pt(10)
        rb.bold = True
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(10)

def table_with_caption(doc, caption, headers, rows):
    # Caption
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_before = Pt(8)
    cp.paragraph_format.space_after = Pt(3)
    cr = cp.add_run(caption)
    cr.font.name = 'Times New Roman'
    cr.font.size = Pt(9)
    cr.font.small_caps = True
    cr.bold = True

    tbl = doc.add_table(rows=1+len(rows), cols=len(headers))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = True

    for i, h in enumerate(headers):
        c = tbl.rows[0].cells[i]
        c.text = ''
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(9)
        r.font.name = 'Times New Roman'
        r.font.color.rgb = RGBColor(255,255,255)
        set_cell_shading(c, "2E4057")
        set_cell_padding(c)

    for ri, rd in enumerate(rows):
        for ci, val in enumerate(rd):
            c = tbl.rows[ri+1].cells[ci]
            c.text = ''
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(val))
            r.font.size = Pt(9)
            r.font.name = 'Times New Roman'
            if ri % 2 == 0:
                set_cell_shading(c, "F0F4F8")
            set_cell_padding(c)

    # Space after table
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(4)
    return tbl

def figure(doc, img_path, caption, width=Inches(5.5)):
    if img_path and os.path.exists(img_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run()
        run.add_picture(img_path, width=width)
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(6)
        r = cap.add_run(caption)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(9)
        r.italic = True

def pseudo_line(doc, line):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = Pt(11)
    p.paragraph_format.left_indent = Cm(1.0)
    r = p.add_run(line)
    r.font.name = 'Consolas'
    r.font.size = Pt(8)


# ══════════════════════════════════════
# MAIN
# ══════════════════════════════════════
def generate_report():
    doc = Document()

    # Page setup - single column, IEEE-ish margins
    for sec in doc.sections:
        sec.top_margin = Cm(2.0)
        sec.bottom_margin = Cm(2.0)
        sec.left_margin = Cm(2.0)
        sec.right_margin = Cm(2.0)

    # Default style
    sty = doc.styles['Normal']
    sty.font.name = 'Times New Roman'
    sty.font.size = Pt(10)
    sty.paragraph_format.space_after = Pt(0)
    sty.paragraph_format.space_before = Pt(0)
    sty.paragraph_format.line_spacing = Pt(13)

    # ══════ TITLE ══════
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("Deep Learning Architectures for Equitable and Accurate Skin Disease Diagnosis")
    r.font.name = 'Times New Roman'
    r.font.size = Pt(20)
    r.bold = True

    # Authors
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Devansh Singh, Prakhar Umrao, Komaravolu Aashrith Saradhi")
    r.font.name = 'Times New Roman'
    r.font.size = Pt(11)

    # Affiliation
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("Department of Computer Science and Engineering, Manipal Institute of Technology, Manipal\nCSE-3223: Deep Learning Lab | Academic Year 2025-26")
    r.font.name = 'Times New Roman'
    r.font.size = Pt(9)
    r.italic = True

    # ── Abstract ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = Pt(12)
    r1 = p.add_run("Abstract\u2014")
    r1.font.name = 'Times New Roman'
    r1.font.size = Pt(9)
    r1.bold = True
    r1.italic = True
    r2 = p.add_run(
        "This project presents an offline-capable, mobile-deployable skin disease classification "
        "system powered by EfficientNet-B2, trained via a 3-stage progressive transfer learning "
        "pipeline across five dermatological datasets (HAM10000, ISIC 2019, PAD-UFES-20, "
        "Fitzpatrick17k-C, DermNet). The system classifies skin lesion images into 16 unified "
        "disease categories spanning neoplastic and inflammatory conditions relevant to Indian "
        "clinical settings. We incorporate Fitzpatrick Skin Type (FST) diversity and employ "
        "experience replay to mitigate catastrophic forgetting during multi-stage training. "
        "The final model achieves 71.1% accuracy on HAM10000 (Stage 1), 67.7% on the combined "
        "ISIC+PAD test set (Stage 2), and demonstrates improved skin tone robustness with a "
        "composite score of 0.364 at Stage 3. The model is exported via PyTorch Mobile (<50 MB) "
        "for fully offline smartphone deployment."
    )
    r2.font.name = 'Times New Roman'
    r2.font.size = Pt(9)

    # Keywords
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = Pt(12)
    rk1 = p.add_run("Keywords\u2014")
    rk1.font.name = 'Times New Roman'
    rk1.font.size = Pt(9)
    rk1.bold = True
    rk1.italic = True
    rk2 = p.add_run("skin disease classification, EfficientNet, transfer learning, experience replay, Fitzpatrick skin type, mobile deployment, dermatological AI")
    rk2.font.name = 'Times New Roman'
    rk2.font.size = Pt(9)
    rk2.italic = True

    # Horizontal rule
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '999999')
    pBdr.append(bottom)
    pPr.append(pBdr)

    # ══════ I. INTRODUCTION ══════
    section_heading(doc, 1, "Introduction")
    body(doc,
        "Skin diseases affect an estimated 1.8 billion people globally [1]. In rural India, "
        "access to dermatological expertise is severely limited\u2014approximately 1 dermatologist "
        "per 100,000 population\u2014leading to delayed diagnoses of malignant and inflammatory "
        "conditions [2]. Early detection of melanoma (MEL), basal cell carcinoma (BCC), and "
        "squamous cell carcinoma (SCC) is critical for patient survival, yet requires "
        "specialized expertise unavailable in primary healthcare settings.",
        indent=False)
    body(doc,
        "Deep learning has shown remarkable promise in automated dermatological diagnosis, "
        "with convolutional neural networks (CNNs) achieving dermatologist-level performance "
        "on curated benchmarks [3]. However, critical challenges remain: (1) the domain gap "
        "between dermoscopic training images and smartphone photographs; (2) demographic bias "
        "from datasets predominantly featuring light-skinned subjects (Fitzpatrick Skin Types "
        "I-II); (3) the need for offline-capable models; and (4) the requirement for a unified "
        "model handling both neoplastic and inflammatory conditions relevant to Indian populations.")
    body(doc,
        "This project addresses these challenges through a 3-stage progressive transfer "
        "learning pipeline using EfficientNet-B2 [4], systematically bridging dermoscopy to "
        "clinical photography. We incorporate Fitzpatrick Skin Type (FST) diversity via the "
        "Fitzpatrick17k-C dataset [5], employ experience replay [6] to mitigate catastrophic "
        "forgetting [7], and produce a compact model deployable via PyTorch Mobile. The "
        "system classifies images into 16 unified disease categories with a safety-first "
        "clinical decision support layer that flags potential malignancies for referral.")

    # ══════ II. LITERATURE REVIEW ══════
    section_heading(doc, 2, "Literature Review")

    subsection_heading(doc, "A", "Deep Learning in Dermatology")
    body(doc,
        "Esteva et al. [3] demonstrated CNN-based skin cancer classification at dermatologist-level "
        "using Inception-v3 on 129,450 images. Subsequent work explored ResNet [8], DenseNet [9], "
        "and EfficientNet [4], with EfficientNet variants achieving state-of-the-art accuracy "
        "through compound scaling of depth, width, and resolution.")

    subsection_heading(doc, "B", "Transfer Learning and Multi-Stage Training")
    body(doc,
        "Transfer learning from ImageNet-pretrained models is standard for medical imaging [10]. "
        "Progressive transfer learning, with sequential adaptation through increasingly specific "
        "domains, improves over direct fine-tuning [11]. Curriculum learning strategies further "
        "enhance convergence in skin lesion classification [12].")

    subsection_heading(doc, "C", "Catastrophic Forgetting and Experience Replay")
    body(doc,
        "Multi-stage training risks catastrophic forgetting, where fine-tuning on new data erases "
        "previously learned representations [7]. Experience replay addresses this by interleaving "
        "samples from previous stages during subsequent training [6], and has been successfully "
        "applied in medical image classification to maintain cross-domain performance.")

    subsection_heading(doc, "D", "Fairness and Skin Tone Bias")
    body(doc,
        "Dermatological AI systems exhibit significant performance disparities across skin tones "
        "[13]. The Fitzpatrick17k dataset [5] provides FST-annotated images across 114 conditions, "
        "enabling bias evaluation. Daneshjou et al. [14] demonstrated up to 20% performance "
        "degradation on darker skin tones in existing classifiers.")

    subsection_heading(doc, "E", "Mobile Deployment and Key Datasets")
    body(doc,
        "Mobile deployment requires model size <50 MB with sub-second inference [15]. Key datasets "
        "include HAM10000 (10,015 dermoscopic images) [17], ISIC 2019 (25,331 images) [18], "
        "PAD-UFES-20 (2,298 smartphone images) [19], Fitzpatrick17k-C, and DermNet. PyTorch Mobile "
        "provides a direct path from research to deployment [16].")

    # ══════ III. RESEARCH GAPS ══════
    section_heading(doc, 3, "Research Gaps")
    body(doc, "This project addresses the following gaps:", indent=False, space_after=2)
    bullet(doc, "Few studies systematically bridge the domain gap from dermoscopy to smartphone photography through progressive training strategies.", bold_prefix="Domain Gap: ")
    bullet(doc, "Training datasets overrepresent FST I-III, creating classifiers with reduced accuracy for darker skin tones prevalent in India.", bold_prefix="Demographic Bias: ")
    bullet(doc, "Existing systems rarely handle neoplastic and inflammatory conditions within a single unified model.", bold_prefix="Unified Taxonomy: ")
    bullet(doc, "Cloud-based AI services are impractical in rural India due to unreliable internet connectivity.", bold_prefix="Offline Deployment: ")
    bullet(doc, "Replay-based continual learning is underexplored in multi-stage dermatological model training.", bold_prefix="Catastrophic Forgetting: ")

    # ══════ IV. OBJECTIVES ══════
    section_heading(doc, 4, "Objectives")
    bullet(doc, "Design a 3-stage progressive transfer learning pipeline adapting dermoscopic models to clinical and smartphone imagery.")
    bullet(doc, "Develop a unified 16-class taxonomy covering neoplastic and inflammatory conditions relevant to Indian populations.")
    bullet(doc, "Incorporate FST diversity via Fitzpatrick17k-C with FST-weighted sampling to reduce performance disparities across skin tones.")
    bullet(doc, "Implement experience replay to mitigate catastrophic forgetting across training stages.")
    bullet(doc, "Deploy on mobile via PyTorch Mobile with model size <50 MB for fully offline operation.")
    bullet(doc, "Integrate a safety-first referral layer that flags potential malignancies for clinical evaluation.")

    # ══════ V. METHODOLOGY ══════
    section_heading(doc, 5, "Proposed Methodology")

    # 5A - Architecture Selection
    subsection_heading(doc, "A", "Architecture Selection")
    body(doc,
        "The mobile deployment constraint (<50 MB) and need for fine-grained lesion detail "
        "guided architecture selection. Table I compares candidate architectures evaluated "
        "for this project.", indent=False)

    table_with_caption(doc, "TABLE I: COMPARISON OF CANDIDATE ARCHITECTURES",
        ["Architecture", "Params (M)", "Input Size", "Size (MB)", "ImageNet Top-1 (%)", "Mobile Fit"],
        [
            ["MobileNetV3-Large", "5.4", "224x224", "~22", "75.2", "Excellent"],
            ["EfficientNet-B0", "5.3", "224x224", "~21", "77.1", "Good"],
            ["EfficientNet-B1", "7.8", "240x240", "~31", "79.1", "Good"],
            ["EfficientNet-B2 *", "9.2", "260x260", "~36", "80.1", "Good"],
            ["EfficientNet-B3", "12.0", "300x300", "~48", "81.6", "Near Limit"],
            ["ResNet-50", "25.6", "224x224", "~98", "76.1", "Too Large"],
            ["DenseNet-121", "8.0", "224x224", "~30", "74.4", "Good"],
            ["ConvNeXt-Tiny", "28.6", "224x224", "~110", "82.1", "Too Large"],
        ]
    )

    body(doc,
        "EfficientNet-B2 (*selected) offers the optimal trade-off: its 260x260 input captures "
        "fine-grained lesion morphology compared to the 224x224 alternatives; 80.1% ImageNet Top-1 "
        "accuracy provides a strong transfer learning foundation (3% higher than B0, 5.7% higher "
        "than DenseNet-121); and the ~36 MB model size fits within the 50 MB budget. EfficientNet-B3 "
        "(48 MB) leaves insufficient headroom for export overhead, while MobileNetV3-Large's 75.2% "
        "accuracy is too low for safety-critical medical screening.")

    # 5B - System Architecture
    subsection_heading(doc, "B", "System Architecture")
    body(doc,
        "The system comprises an EfficientNet-B2 backbone (ImageNet-pretrained, 1,408-dimensional "
        "feature vectors) with a custom classification head: AdaptiveAvgPool2d(1) followed by "
        "BatchNorm1d(1408), Dropout(p=0.3), and Linear(1408, 16). Fig. 1 illustrates the "
        "complete architecture.", indent=False)
    figure(doc, IMG_ARCH, "Fig. 1. EfficientNet-B2 system architecture with 16-class output head.")

    # 5C - Unified Taxonomy
    subsection_heading(doc, "C", "Unified 16-Class Taxonomy")
    body(doc,
        "Five datasets with diverse label schemes are harmonized into 16 clinically meaningful "
        "categories (Table II), divided into neoplastic (9 classes covering malignant, pre-malignant, "
        "and benign lesions) and inflammatory (7 classes covering conditions prevalent in Indian "
        "clinical settings).", indent=False)

    table_with_caption(doc, "TABLE II: UNIFIED 16-CLASS DISEASE TAXONOMY",
        ["Code", "Full Name", "Category"],
        [
            ["MEL", "Melanoma", "Malignant"],
            ["BCC", "Basal Cell Carcinoma", "Malignant"],
            ["SCC", "Squamous Cell Carcinoma", "Malignant"],
            ["AK", "Actinic Keratosis", "Pre-malignant"],
            ["NEV", "Melanocytic Nevus", "Benign"],
            ["BKL", "Benign Keratosis", "Benign"],
            ["DF", "Dermatofibroma", "Benign"],
            ["VASC", "Vascular Lesion", "Benign"],
            ["SEK", "Seborrheic Keratosis", "Benign"],
            ["TINEA", "Tinea (Ringworm)", "Inflammatory"],
            ["PSORIASIS", "Psoriasis", "Inflammatory"],
            ["VITILIGO", "Vitiligo", "Inflammatory"],
            ["MELASMA", "Melasma", "Inflammatory"],
            ["FUNGAL", "Fungal Infection", "Inflammatory"],
            ["ECZEMA", "Eczema / Dermatitis", "Inflammatory"],
            ["URTICARIA", "Urticaria (Hives)", "Inflammatory"],
        ]
    )

    # 5D - Data Preprocessing
    subsection_heading(doc, "D", "Data Preprocessing Pipeline")
    body(doc,
        "The preprocessing pipeline harmonizes five heterogeneous datasets through: (1) label "
        "mapping to the unified taxonomy (e.g., Fitzpatrick17k's 114 diagnoses mapped to 16 "
        "classes via keyword matching); (2) MD5-based deduplication of ISIC 2019 against HAM10000 "
        "(removing 7,054 duplicate images); (3) patient-level GroupShuffleSplit (70/15/15 ratio) "
        "preventing data leakage; and (4) image validation to remove corrupt files. Table III "
        "summarizes dataset configurations. Fig. 2 illustrates the preprocessing flow.", indent=False)

    table_with_caption(doc, "TABLE III: DATASET CONFIGURATION SUMMARY",
        ["Dataset", "Total", "Type", "Train", "Val", "Test"],
        [
            ["HAM10000", "10,015", "Dermoscopy", "7,053", "1,465", "1,497"],
            ["ISIC 2019 (deduped)", "18,277", "Dermoscopy", "13,255", "2,494", "2,528"],
            ["PAD-UFES-20", "2,298", "Smartphone", "1,582", "370", "346"],
            ["Fitzpatrick17k-C", "2,553*", "Clinical Atlas", "2,553", "352", "\u2014"],
            ["DermNet", "10,390", "Clinical Atlas", "8,831", "1,559", "\u2014"],
            ["Stage 3 Replay", "26,065", "Mixed", "26,065", "4,729", "6,163"],
        ]
    )
    p = doc.add_paragraph()
    r = p.add_run("*11,394 original images filtered to 2,553 after mapping 114 diagnoses to 16 unified classes.")
    r.font.size = Pt(8)
    r.font.name = 'Times New Roman'
    r.italic = True

    figure(doc, IMG_PREPROCESS, "Fig. 2. Data preprocessing pipeline across five heterogeneous datasets.")

    # 5E - Training Pipeline
    subsection_heading(doc, "E", "Three-Stage Progressive Transfer Learning")
    body(doc,
        "The model is progressively trained through three stages, each building upon the "
        "previous checkpoint. Fig. 3 illustrates the complete pipeline.", indent=False)
    figure(doc, IMG_PIPELINE, "Fig. 3. Three-stage progressive transfer learning pipeline.")

    # Stage 1
    body(doc, "Stage 1 \u2014 Dermoscopy Baseline (HAM10000):", indent=False, space_after=2, space_before=6)
    body(doc,
        "EfficientNet-B2 initialized with ImageNet weights. Backbone frozen except last 3 blocks "
        "and classification head. Trained for 25 epochs using Adam optimizer (lr=1e-4, "
        "weight_decay=1e-4) with CosineAnnealingLR (T_max=25, eta_min=1e-6). Loss: CrossEntropyLoss "
        "with inverse-frequency class weights. WeightedRandomSampler addresses class imbalance. "
        "Data augmentation includes horizontal/vertical flips, rotation(20\u00b0), ColorJitter "
        "(brightness=0.3, contrast=0.3, saturation=0.25, hue=0.1), and RandomErasing(p=0.2). "
        "Checkpoint: stage1.pth (best validation loss).")

    # Stage 2
    body(doc, "Stage 2 \u2014 Domain Bridge (ISIC 2019 + PAD-UFES-20):", indent=False, space_after=2, space_before=6)
    body(doc,
        "Stage 1 checkpoint loaded; 7-class head expanded to 16 classes by copying overlapping "
        "weights and Xavier-initializing new class weights. ISIC 2019 deduplicated against "
        "HAM10000 via MD5 hashes (7,054 removed). Last 5 blocks + head unfrozen with differential "
        "learning rates (backbone: 1e-5, head: 1e-4). Trained for 20 epochs with stronger "
        "augmentation (ColorJitter brightness=0.5, contrast=0.5, saturation=0.4, hue=0.15) to "
        "simulate smartphone camera variation. Checkpoint: stage2.pth (best validation loss, epoch 19).")

    # Stage 3
    body(doc, "Stage 3 \u2014 Skin Tone Robustness with Experience Replay:", indent=False, space_after=2, space_before=6)
    body(doc,
        "Stage 2 checkpoint loaded; all parameters unfrozen. Experience replay interleaves "
        "Stage 2 training data with Fitzpatrick17k-C and DermNet. A WeightedRandomSampler with "
        "joint class \u00d7 FST-band weighting oversamples underrepresented skin tones. Source-aware "
        "label smoothing (DermNet: 0.12, Fitz/ISIC/PAD: 0.04) accounts for dataset noise. "
        "Trained for 12 epochs with Adam (backbone: 5e-6, head: 5e-5) and CosineAnnealingLR "
        "(T_max=12, eta_min=1e-7). Best checkpoint selected using the composite scoring "
        "function (Eq. 1). The final FitzTuned checkpoint is produced by further fine-tuning "
        "on Fitzpatrick data.")

    # 5F - Composite Score
    subsection_heading(doc, "F", "Composite Checkpoint Selection")
    body(doc, "The composite score balances fairness and safety:", indent=False, space_after=2)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("S = 0.35\u00b7F1_PAD + 0.25\u00b7R_mal + 0.20\u00b7F1_dark + 0.20\u00b7BA   (1)")
    r.font.name = 'Times New Roman'
    r.font.size = Pt(10)
    r.italic = True
    body(doc,
        "where F1_PAD is PAD-UFES-20 macro F1, R_mal is mean recall of malignant classes "
        "(MEL, SCC, AK), F1_dark is F1 on dark skin (FST V-VI), and BA is balanced accuracy.")

    # 5G - Pseudocode
    subsection_heading(doc, "G", "Training Algorithm")
    pseudo_lines = [
        "Algorithm 1: Three-Stage Progressive Transfer Learning",
        "\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500\u2500",
        "Input: HAM10000, ISIC2019, PAD, Fitz17k-C, DermNet",
        "Output: stage3_replay_fitz_tuned.pth",
        "",
        "1. STAGE 1 (Dermoscopy Baseline):",
        "   model <- EffNetB2(pretrained=ImageNet, classes=16)",
        "   FREEZE backbone; UNFREEZE last_3_blocks + head",
        "   FOR epoch = 1..25: TRAIN(HAM10000, WeightedSampler)",
        "   SAVE best(val_loss) -> stage1.pth",
        "",
        "2. STAGE 2 (Domain Bridge):",
        "   model <- LOAD(stage1.pth)",
        "   head <- EXPAND(7->16, copy overlapping weights)",
        "   ISIC2019 <- DEDUP(ISIC2019, HAM, MD5)  // -7054",
        "   UNFREEZE last_5_blocks + head",
        "   FOR epoch = 1..20: TRAIN(ISIC+PAD, diff_lr)",
        "   SAVE best(val_loss) -> stage2.pth",
        "",
        "3. STAGE 3 (Replay + Skin Tone):",
        "   model <- LOAD(stage2.pth); UNFREEZE ALL",
        "   replay <- CONCAT(S2_train, Fitz, DermNet)",
        "   sampler <- Weight(class_freq * fst_band)",
        "   FOR epoch = 1..12:",
        "     TRAIN(replay, source_label_smoothing)",
        "     score <- 0.35*PAD_f1 + 0.25*mal_recall",
        "           + 0.20*dark_f1 + 0.20*bal_acc",
        "   SAVE best(score) -> stage3_replay_best.pth",
        "   FINETUNE(Fitz) -> stage3_replay_fitz_tuned.pth",
    ]
    for line in pseudo_lines:
        pseudo_line(doc, line)

    # Small spacer
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # 5H - Inference Pipeline
    subsection_heading(doc, "H", "Mobile Inference Pipeline")
    body(doc,
        "The trained model is exported via PyTorch Mobile's Lite Interpreter (.ptl, <50 MB). "
        "An InferenceWrapper produces softmax probabilities, confidence score, and top-3 "
        "predicted class indices. Pre-inference quality checks reject blurry images (Laplacian "
        "variance < 100) and poorly lit images (mean brightness outside 40\u2013220). A safety-first "
        "layer flags any malignant class (MEL, BCC, SCC, AK) appearing in the top-3 for "
        "immediate clinic referral. Fig. 4 illustrates the deployment pipeline.", indent=False)
    figure(doc, IMG_INFERENCE, "Fig. 4. Mobile inference and clinical decision support pipeline.")

    # ══════ VI. EXPERIMENTAL RESULTS ══════
    section_heading(doc, 6, "Experimental Results")

    # 6A - Stage 1
    subsection_heading(doc, "A", "Stage 1: HAM10000 Baseline")
    body(doc,
        "The Stage 1 model trained for 25 epochs achieves strong dermoscopic classification. "
        "Table IV presents the best checkpoint performance on the HAM10000 test set.", indent=False)

    table_with_caption(doc, "TABLE IV: STAGE 1 \u2014 HAM10000 TEST SET RESULTS (BEST CHECKPOINT, EPOCH 25)",
        ["Metric", "Value"],
        [
            ["Accuracy", "0.711"],
            ["Macro F1-Score", "0.676"],
            ["Balanced Accuracy", "0.755"],
        ]
    )

    # 6B - Stage 2
    subsection_heading(doc, "B", "Stage 2: ISIC 2019 + PAD-UFES-20")
    body(doc,
        "Stage 2 bridges to smartphone imagery. The best checkpoint (epoch 19, by validation "
        "loss = 0.861) is evaluated on both the PAD smartphone test set and the combined test "
        "set. Tables V and VI present the results.", indent=False)

    table_with_caption(doc, "TABLE V: STAGE 2 \u2014 TEST SET METRICS (BEST CHECKPOINT, EPOCH 19)",
        ["Test Set", "Images", "Accuracy", "Macro F1", "Balanced Acc"],
        [
            ["PAD-UFES-20 (smartphone)", "346", "0.668", "0.416", "0.547"],
            ["Combined ISIC + PAD", "2,874", "0.677", "0.564", "0.531"],
        ]
    )

    table_with_caption(doc, "TABLE VI: STAGE 2 \u2014 PER-CLASS REPORT (COMBINED TEST SET, 2,874 IMAGES)",
        ["Class", "Precision", "Recall", "F1-Score", "Support"],
        [
            ["MEL", "0.73", "0.58", "0.64", "618"],
            ["BCC", "0.71", "0.74", "0.72", "590"],
            ["SCC", "0.37", "0.18", "0.24", "131"],
            ["AK", "0.52", "0.44", "0.47", "222"],
            ["NEV", "0.71", "0.90", "0.80", "962"],
            ["BKL", "0.49", "0.44", "0.46", "262"],
            ["DF", "0.45", "0.26", "0.33", "19"],
            ["VASC", "0.93", "0.56", "0.70", "25"],
            ["SEK", "0.72", "0.69", "0.70", "45"],
        ]
    )

    # 6C - Stage 3 Original
    subsection_heading(doc, "C", "Stage 3 Original: Without Replay (Catastrophic Forgetting)")
    body(doc,
        "The initial Stage 3 model trained on Fitzpatrick17k-C + DermNet without replay exhibits "
        "severe catastrophic forgetting: SCC and AK recall collapse to 0.000 as dermoscopic "
        "knowledge is overwritten by clinical atlas features. Table VII shows evaluation on the "
        "Fitzpatrick+DermNet val/test split. This motivates the replay approach.", indent=False)

    table_with_caption(doc, "TABLE VII: STAGE 3 ORIGINAL \u2014 FITZ+DERMNET EVALUATION (3,683 IMAGES, BEST EPOCH 12)",
        ["Metric", "Value"],
        [
            ["Accuracy", "0.518"],
            ["Macro F1-Score", "0.380"],
            ["Balanced Accuracy", "0.413"],
            ["MEL Recall", "0.543"],
            ["SCC Recall", "0.000 (forgetting)"],
            ["AK Recall", "0.000 (forgetting)"],
            ["TINEA Recall", "0.660 (new)"],
            ["FUNGAL Recall", "0.870 (new)"],
            ["ECZEMA Recall", "0.470 (new)"],
        ]
    )

    # 6D - Stage 3 Replay
    subsection_heading(doc, "D", "Stage 3 Replay: With Experience Replay")
    body(doc,
        "Replay training interleaves Stage 2 data with Fitzpatrick17k and DermNet using "
        "FST-weighted sampling and source-aware label smoothing. The best checkpoint is "
        "selected at epoch 7 (composite score = 0.364). Table VIII shows the key metrics, "
        "demonstrating successful recovery of SCC recall from 0.000 to 0.458.", indent=False)

    table_with_caption(doc, "TABLE VIII: STAGE 3 REPLAY \u2014 BEST CHECKPOINT (EPOCH 7, COMPOSITE SCORE = 0.364)",
        ["Metric", "Value"],
        [
            ["Train Macro F1", "0.580"],
            ["Validation Macro F1", "0.419"],
            ["PAD-UFES-20 Monitor F1", "0.356"],
            ["PAD Balanced Accuracy", "0.486"],
            ["MEL Recall (PAD)", "0.250"],
            ["SCC Recall (PAD)", "0.458 (recovered)"],
            ["AK Recall (PAD)", "0.755 (maintained)"],
        ]
    )

    # 6E - Cross-Stage
    subsection_heading(doc, "E", "Cross-Stage Performance Summary")
    body(doc,
        "Table IX summarizes results across all training stages, demonstrating the progressive "
        "improvement and the impact of replay on preventing catastrophic forgetting.", indent=False)

    table_with_caption(doc, "TABLE IX: CROSS-STAGE PERFORMANCE COMPARISON",
        ["Checkpoint", "Evaluation Set", "Accuracy", "Macro F1", "Key Observation"],
        [
            ["Stage 1", "HAM10000 test (1,497)", "0.711", "0.676", "Strong dermoscopy baseline"],
            ["Stage 2", "ISIC+PAD combined (2,874)", "0.677", "0.564", "Smartphone domain acquired"],
            ["Stage 3 Orig.", "Fitz+DermNet val/test (3,683)", "0.518", "0.380", "SCC/AK=0 (forgetting)"],
            ["Stage 3 Replay", "PAD monitor (346)", "0.486 (BA)", "0.356", "SCC/AK recovered"],
            ["FitzTuned", "Replay val (4,729)", "0.519 (BA)", "0.436", "Best overall model"],
        ]
    )

    # ══════ VII. CONCLUSION ══════
    section_heading(doc, 7, "Conclusion and Future Work")

    subsection_heading(doc, "A", "Conclusion")
    body(doc,
        "We presented a 3-stage progressive transfer learning approach for equitable skin "
        "disease classification, demonstrating: (1) successful domain bridging from dermoscopy "
        "to smartphone imagery across HAM10000, ISIC 2019, and PAD-UFES-20; (2) experience "
        "replay effectively mitigating catastrophic forgetting (SCC recall from 0.000 without "
        "replay to 0.458 with replay); (3) FST-weighted training addressing demographic bias "
        "across skin tones; (4) a compact <50 MB model enabling fully offline mobile deployment; "
        "and (5) a safety-first clinical decision support layer that flags malignancies for "
        "clinic referral.", indent=False)

    subsection_heading(doc, "B", "Limitations")
    body(doc,
        "Inflammatory class performance is limited by the absence of smartphone training data "
        "for these conditions. SCC and AK recall (0.458, 0.755) remain below clinically "
        "desirable thresholds (>0.80) for safety-critical screening. The model has not been "
        "validated on real-world images from the target deployment population in rural India.", indent=False)

    subsection_heading(doc, "C", "Future Work")
    bullet(doc, "Integrate SCIN dataset (10,000+ smartphone images with FST labels) for inflammatory condition coverage.")
    bullet(doc, "Collect 300-500 images from Karnataka clinics for target-population validation.")
    bullet(doc, "Implement Grad-CAM visualization for explainable predictions.")
    bullet(doc, "Explore two-model architecture (neoplastic + inflammatory classifiers).")
    bullet(doc, "Expand taxonomy: ACNE, SCABIES, INSECT_BITE, WOUND_INFECTION.")

    # ══════ VIII. REFERENCES ══════
    section_heading(doc, 8, "References")
    refs = [
        '[1] R. Hay et al., "The global burden of skin disease in 2010," J. Invest. Dermatol., vol. 134, no. 6, pp. 1527-1534, 2014.',
        '[2] S. Srinivasan et al., "Dermatology workforce and service delivery in India," Indian J. Dermatol., vol. 64, pp. 1-4, 2019.',
        '[3] A. Esteva et al., "Dermatologist-level classification of skin cancer with deep neural networks," Nature, vol. 542, pp. 115-118, 2017.',
        '[4] M. Tan and Q. Le, "EfficientNet: Rethinking model scaling for CNNs," ICML, pp. 6105-6114, 2019.',
        '[5] M. Groh et al., "Evaluating DNNs trained on clinical images with the Fitzpatrick 17k dataset," CVPR, pp. 1820-1828, 2021.',
        '[6] D. Rolnick et al., "Experience replay for continual learning," NeurIPS, vol. 32, pp. 350-360, 2019.',
        '[7] M. McCloskey and N. Cohen, "Catastrophic interference in connectionist networks," Psychol. Learn. Motiv., vol. 24, pp. 109-165, 1989.',
        '[8] K. He et al., "Deep residual learning for image recognition," CVPR, pp. 770-778, 2016.',
        '[9] G. Huang et al., "Densely connected convolutional networks," CVPR, pp. 4700-4708, 2017.',
        '[10] S. Pan and Q. Yang, "A survey on transfer learning," IEEE TKDE, vol. 22, no. 10, pp. 1345-1359, 2010.',
        '[11] J. Yosinski et al., "How transferable are features in deep neural networks?," NeurIPS, vol. 27, pp. 3320-3328, 2014.',
        '[12] Y. Bengio et al., "Curriculum learning," ICML, pp. 41-48, 2009.',
        '[13] A. Adamson and A. Smith, "ML and health care disparities in dermatology," JAMA Dermatol., vol. 154, pp. 1247-1248, 2018.',
        '[14] R. Daneshjou et al., "Disparities in dermatology AI performance," Sci. Adv., vol. 8, 2022.',
        '[15] Y. Chen et al., "Deep learning on mobile and embedded devices," ACM Comput. Surv., vol. 53, pp. 1-37, 2020.',
        '[16] PyTorch Mobile Documentation, pytorch.org/mobile, 2023.',
        '[17] P. Tschandl et al., "The HAM10000 dataset," Sci. Data, vol. 5, pp. 1-9, 2018.',
        '[18] N. Codella et al., "Skin lesion analysis toward melanoma detection 2018: ISIC challenge," arXiv:1902.03368, 2019.',
        '[19] A. Pacheco et al., "PAD-UFES-20: Skin lesion dataset from smartphones," Data Brief, vol. 32, 2020.',
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = Pt(11)
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.first_line_indent = Cm(-0.5)
        r = p.add_run(ref)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(9)

    # ── Save ──
    doc.save(str(OUTPUT_PATH))
    print(f"\n[OK] Report saved to: {OUTPUT_PATH}")
    print(f"   File size: {OUTPUT_PATH.stat().st_size / 1024:.1f} KB")


if __name__ == '__main__':
    generate_report()
